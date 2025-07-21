from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def add_border_to_paragraph(paragraph):
    """Menambahkan border ke paragraph"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pBdr.set(qn('w:top'), qn('w:single'))
    pBdr.set(qn('w:left'), qn('w:single'))
    pBdr.set(qn('w:bottom'), qn('w:single'))
    pBdr.set(qn('w:right'), qn('w:single'))
    pPr.append(pBdr)

def create_form_evaluasi():
    # Membuat dokumen baru
    doc = Document()
    
    # Set margin dokumen
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    # Header dokumen
    header = doc.add_heading('FORM EVALUASI KINERJA GURU', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header.runs[0]
    header_run.font.size = Pt(16)
    header_run.font.name = 'Arial'
    header_run.bold = True
    
    # Sub header
    subheader = doc.add_paragraph('SMP PENIDA KATAPANG')
    subheader.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subheader_run = subheader.runs[0]
    subheader_run.font.size = Pt(14)
    subheader_run.font.name = 'Arial'
    subheader_run.bold = True
    
    # Tahun akademik
    tahun = doc.add_paragraph('TAHUN AKADEMIK 2024/2025')
    tahun.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tahun_run = tahun.runs[0]
    tahun_run.font.size = Pt(12)
    tahun_run.font.name = 'Arial'
    
    # Line break
    doc.add_paragraph()
    
    # Informasi evaluator (sudah diisi)
    info_table = doc.add_table(rows=4, cols=4)
    info_table.style = 'Table Grid'
    info_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    # Data evaluator
    evaluator_data = [
        ['Nama Evaluator', ':', 'Dr. H. Asep Suryadi, M.Pd', ''],
        ['Jabatan', ':', 'Kepala Sekolah', ''],
        ['Tanggal Evaluasi', ':', '15 Januari 2025', ''],
        ['Periode Evaluasi', ':', 'Semester 1 / Tahun 2024/2025', '']
    ]
    
    for i, row_data in enumerate(evaluator_data):
        for j, cell_data in enumerate(row_data):
            cell = info_table.cell(i, j)
            cell.text = cell_data
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # Bold untuk data yang sudah diisi
            if j == 2 and cell_data:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            if j == 0:
                cell.width = Inches(2)
            elif j == 1:
                cell.width = Inches(0.3)
            else:
                cell.width = Inches(2.5)
    
    doc.add_paragraph()
    
    # GURU 1 - Budi Santoso (Ranking tertinggi)
    guru1_header = doc.add_heading('EVALUASI GURU 1', level=1)
    guru1_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    guru1_run = guru1_header.runs[0]
    guru1_run.font.size = Pt(14)
    guru1_run.font.name = 'Arial'
    
    # Data Guru 1 (sudah diisi)
    guru1_table = doc.add_table(rows=4, cols=4)
    guru1_table.style = 'Table Grid'
    
    guru1_data = [
        ['Nama Guru', ':', 'Budi Santoso, S.Pd', ''],
        ['NIP/NUPTK', ':', '196805121990031005', ''],
        ['Mata Pelajaran', ':', 'IPA (Ilmu Pengetahuan Alam)', ''],
        ['Kelas yang Diampu', ':', 'VII A, VII B, VIII A', '']
    ]
    
    for i, row_data in enumerate(guru1_data):
        for j, cell_data in enumerate(row_data):
            cell = guru1_table.cell(i, j)
            cell.text = cell_data
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # Bold untuk data yang sudah diisi
            if j == 2 and cell_data:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            if j == 0:
                cell.width = Inches(2)
            elif j == 1:
                cell.width = Inches(0.3)
            else:
                cell.width = Inches(2.5)
    
    # Kriteria Evaluasi Guru 1
    doc.add_paragraph()
    kriteria_header1 = doc.add_paragraph('KRITERIA EVALUASI:')
    kriteria_run1 = kriteria_header1.runs[0]
    kriteria_run1.font.size = Pt(12)
    kriteria_run1.font.name = 'Arial'
    kriteria_run1.bold = True
    
    # Tabel evaluasi Guru 1 dengan data terisi
    eval_table1 = doc.add_table(rows=7, cols=6)
    eval_table1.style = 'Table Grid'
    eval_table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header tabel evaluasi
    headers = ['No', 'Kriteria Evaluasi', 'Sangat Baik (90-100)', 'Baik (80-89)', 'Cukup (70-79)', 'Kurang (<70)']
    for j, header in enumerate(headers):
        cell = eval_table1.cell(0, j)
        cell.text = header
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Data kriteria evaluasi Guru 1 (Budi Santoso) - sudah terisi berdasarkan nilai Excel
    kriteria_eval1 = [
        ['1', 'Kedisiplinan\n(Ketepatan waktu, kehadiran)', '□', '☑', '□', '□'],  # 80
        ['2', 'Penguasaan Materi\n(Kemampuan mengajar)', '☑', '□', '□', '□'],    # 95
        ['3', 'Metode Mengajar\n(Variasi dan kreativitas)', '☑', '□', '□', '□'], # 90
        ['4', 'Komunikasi\n(Interaksi dengan siswa)', '□', '☑', '□', '□'],       # 85
        ['5', 'Evaluasi Pembelajaran\n(Sistem penilaian)', '□', '☑', '□', '□'],  # 80
        ['', 'TOTAL NILAI AHP', '', '', '87.0', '']
    ]
    
    for i, row_data in enumerate(kriteria_eval1, 1):
        for j, cell_data in enumerate(row_data):
            cell = eval_table1.cell(i, j)
            cell.text = cell_data
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                if j == 0 or j > 1:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Bold untuk total dan highlight nilai
            if i == 6 or (j == 4 and i == 6):
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        if j == 4:  # Nilai AHP
                            run.font.size = Pt(12)
    
    # Komentar Guru 1 (sudah diisi)
    doc.add_paragraph()
    comment1 = doc.add_paragraph('KOMENTAR DAN SARAN:')
    comment1_run = comment1.runs[0]
    comment1_run.font.size = Pt(12)
    comment1_run.font.name = 'Arial'
    comment1_run.bold = True
    
    # Komentar terisi
    comment_text1 = [
        "Guru Budi Santoso menunjukkan kinerja yang sangat baik dengan nilai AHP 87.0.",
        "Keunggulan: Penguasaan materi sangat baik dan metode mengajar yang inovatif.",
        "Rekomendasi: Pertahankan kinerja dan jadikan mentor untuk guru lain.",
        "Perlu sedikit perbaikan dalam hal kedisiplinan dan evaluasi pembelajaran."
    ]
    
    for text in comment_text1:
        line = doc.add_paragraph(text)
        line_run = line.runs[0]
        line_run.font.size = Pt(10)
    
    # Page break
    doc.add_page_break()
    
    # GURU 2 - Siti Nurhaliza (Ranking kedua)
    guru2_header = doc.add_heading('EVALUASI GURU 2', level=1)
    guru2_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    guru2_run = guru2_header.runs[0]
    guru2_run.font.size = Pt(14)
    guru2_run.font.name = 'Arial'
    
    # Data Guru 2 (sudah diisi)
    guru2_table = doc.add_table(rows=4, cols=4)
    guru2_table.style = 'Table Grid'
    
    guru2_data = [
        ['Nama Guru', ':', 'Siti Nurhaliza, S.Pd', ''],
        ['NIP/NUPTK', ':', '197203151998022003', ''],
        ['Mata Pelajaran', ':', 'Bahasa Indonesia', ''],
        ['Kelas yang Diampu', ':', 'VII C, VIII B, IX A', '']
    ]
    
    for i, row_data in enumerate(guru2_data):
        for j, cell_data in enumerate(row_data):
            cell = guru2_table.cell(i, j)
            cell.text = cell_data
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # Bold untuk data yang sudah diisi
            if j == 2 and cell_data:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            if j == 0:
                cell.width = Inches(2)
            elif j == 1:
                cell.width = Inches(0.3)
            else:
                cell.width = Inches(2.5)
    
    # Kriteria Evaluasi Guru 2
    doc.add_paragraph()
    kriteria_header2 = doc.add_paragraph('KRITERIA EVALUASI:')
    kriteria_run2 = kriteria_header2.runs[0]
    kriteria_run2.font.size = Pt(12)
    kriteria_run2.font.name = 'Arial'
    kriteria_run2.bold = True
    
    # Tabel evaluasi Guru 2 dengan data terisi
    eval_table2 = doc.add_table(rows=7, cols=6)
    eval_table2.style = 'Table Grid'
    eval_table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header tabel evaluasi Guru 2
    for j, header in enumerate(headers):
        cell = eval_table2.cell(0, j)
        cell.text = header
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Data kriteria evaluasi Guru 2 (Siti Nurhaliza) - sudah terisi berdasarkan nilai Excel
    kriteria_eval2 = [
        ['1', 'Kedisiplinan\n(Ketepatan waktu, kehadiran)', '☑', '□', '□', '□'],  # 90
        ['2', 'Penguasaan Materi\n(Kemampuan mengajar)', '□', '☑', '□', '□'],    # 85
        ['3', 'Metode Mengajar\n(Variasi dan kreativitas)', '□', '☑', '□', '□'], # 85
        ['4', 'Komunikasi\n(Interaksi dengan siswa)', '□', '☑', '□', '□'],       # 80
        ['5', 'Evaluasi Pembelajaran\n(Sistem penilaian)', '☑', '□', '□', '□'],  # 90
        ['', 'TOTAL NILAI AHP', '', '', '86.0', '']
    ]
    
    for i, row_data in enumerate(kriteria_eval2, 1):
        for j, cell_data in enumerate(row_data):
            cell = eval_table2.cell(i, j)
            cell.text = cell_data
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                if j == 0 or j > 1:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Bold untuk total dan highlight nilai
            if i == 6 or (j == 4 and i == 6):
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        if j == 4:  # Nilai AHP
                            run.font.size = Pt(12)
    
    # Komentar Guru 2 (sudah diisi)
    doc.add_paragraph()
    comment2 = doc.add_paragraph('KOMENTAR DAN SARAN:')
    comment2_run = comment2.runs[0]
    comment2_run.font.size = Pt(12)
    comment2_run.font.name = 'Arial'
    comment2_run.bold = True
    
    # Komentar terisi
    comment_text2 = [
        "Guru Siti Nurhaliza menunjukkan kinerja yang baik dengan nilai AHP 86.0.",
        "Keunggulan: Kedisiplinan sangat baik dan sistem evaluasi pembelajaran yang efektif.",
        "Rekomendasi: Tingkatkan variasi metode mengajar dan penguasaan materi.",
        "Komunikasi dengan siswa perlu lebih ditingkatkan untuk hasil yang optimal."
    ]
    
    for text in comment_text2:
        line = doc.add_paragraph(text)
        line_run = line.runs[0]
        line_run.font.size = Pt(10)
    
    # Footer dengan tanda tangan
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Tabel tanda tangan (sudah diisi)
    signature_table = doc.add_table(rows=4, cols=2)
    signature_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    signature_data = [
        ['Evaluator', 'Kepala Sekolah'],
        ['', ''],
        ['', ''],
        ['(Dr. H. Asep Suryadi, M.Pd)', '(Dr. H. Asep Suryadi, M.Pd)']
    ]
    
    for i, row_data in enumerate(signature_data):
        for j, cell_data in enumerate(row_data):
            cell = signature_table.cell(i, j)
            cell.text = cell_data
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    if i == 0 or i == 3:
                        run.font.bold = True
    
    # Save dokumen
    doc.save('/Users/flashcode/Documents/project-destra/Form_Evaluasi_Guru_Terisi_SMP_PENIDA_KATAPANG_FIX.docx')
    print('Form evaluasi guru (sudah terisi) berhasil dibuat!')
    print('File tersimpan di: /Users/flashcode/Documents/project-destra/Form_Evaluasi_Guru_Terisi_SMP_PENIDA_KATAPANG_FIX.docx')

if __name__ == "__main__":
    create_form_evaluasi()